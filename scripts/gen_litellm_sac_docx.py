#!/usr/bin/env python3
"""Generate 4 SAC DOCX reports for LiteLLM HK:
  1. 部署指南 (EN) — renamed from "技术交付报告"
  2. 部署指南 (ZH)
  3. Business Value Report (EN)
  4. Business Value Report (ZH)

  Structure follows Huawei Cloud support doc pattern:
    hermes_01=方案概述 → 02=资源成本 → 03=实施步骤(父) →
    04=准备工作 → 05=快速部署 → 06=开始使用 → 07=快速卸载 →
    08=附录 → 09=修订记录
"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

BASE = r'C:\Users\Administrator\Desktop\Project\claudeproject\solution-implementations'
OUT = os.path.join(BASE, 'release', 'litellm', 'hk')


def set_cell_shading(cell, color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_styled_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '2F5496')
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, 'D6E4F0')
    return table


def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.0
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2"/>')
    p._element.get_or_add_pPr().append(shading)
    return p


# =============================================================
#  1. 部署指南 (EN) - Hermes 9-chapter structure
# =============================================================
def gen_deploy_en():
    doc = Document()

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(120)
    run = p.add_run('Deployment Guide'); run.bold = True; run.font.size = Pt(26); run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('LiteLLM, The Unified LLM Gateway'); run.font.size = Pt(16); run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

    for line in ['', '', 'Document Type: SAC Deployment Guide',
                 'Solution Version: v1.0', 'Release Date: 2026-06-10', 'Region: Hong Kong (ap-southeast-1)']:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line); run.font.size = Pt(11); run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_page_break()

    # === Chapter 1: 方案概述 (Solution Overview) ===
    doc.add_heading('1. Solution Overview', level=1)

    doc.add_heading('1.1 Introduction', level=2)
    doc.add_paragraph(
        'This solution deploys LiteLLM (GitHub 17k+ Stars) — the unified LLM API gateway — on a '
        'Huawei Cloud Flexus X instance via a single Terraform template. LiteLLM unifies 100+ LLM '
        'providers (OpenAI, Anthropic, Azure, Cohere, etc.) into a single OpenAI-compatible API, '
        'with built-in virtual key management, spend tracking, load balancing, and an Admin Web UI.')

    doc.add_heading('1.2 Architecture', level=2)
    doc.add_paragraph('Three containers orchestrated via Docker Compose:')
    add_styled_table(doc,
        ['Container', 'Image', 'Role'],
        [['LiteLLM Proxy', 'ghcr.io/berriai/litellm-database:main-stable', 'API gateway, routing, key mgmt'],
         ['PostgreSQL 16', 'postgres:16', 'Virtual keys, spend data, config'],
         ['Prometheus', 'prom/prometheus:v2.53.0', 'Metrics & monitoring']])

    doc.add_paragraph('Data flow: Client → LiteLLM Proxy (:4000) → key verification (PG) → provider API → response + spend record')
    add_code_block(doc,
        '┌──────────┐   POST /v1/chat/completions   ┌────────────┐   route   ┌──────────┐\n'
        '│  Client   │ ──────────────────────────►  │  LiteLLM    │ ──────►  │ OpenAI   │\n'
        '│  App      │ ◄──────────────────────────  │  Proxy:4000 │ ◄──────  │ Anthropic│\n'
        '└──────────┘       response                └──────┬─────┘          │ Azure    │\n'
        '                                                   │ DB verify      └──────────┘\n'
        '                                          ┌───────▼──────┐\n'
        '                                          │ PostgreSQL   │\n'
        '                                          │  Prometheus  │\n'
        '                                          └──────────────┘')

    doc.add_heading('1.3 Key Benefits', level=2)
    for b in ['100+ LLM providers unified under one OpenAI-compatible API',
              'Virtual key management with per-key budget limits and spend tracking',
              'Automatic failover and load balancing across providers',
              'One-click deploy via RFS, ~10 minutes to production',
              'Production-ready — 8vCPU 16GB Flexus X with performance mode, 200GB SSD']:
        doc.add_paragraph(b, style='List Bullet')

    doc.add_heading('1.4 Constraints & Limitations', level=2)
    for c in ['Huawei Cloud account with sufficient balance required',
              'RFS (Resource Formation Service) must be enabled',
              'IAM sub-users need rf_admin_trust delegation',
              'LLM API costs billed separately by each provider']:
        doc.add_paragraph(c, style='List Bullet')

    # === Chapter 2: 资源成本 (Resource & Cost Planning) ===
    doc.add_heading('2. Resource and Cost Planning', level=1)
    add_styled_table(doc,
        ['Service', 'Specification', 'Qty', 'Monthly (USD)'],
        [['Flexus X Instance', 'x1.4u.8g (8vCPU 16GB), Ubuntu 24.04, performance mode', '1', '~$41'],
         ['System Disk', '200GB high-IO SSD', '1', '~$14'],
         ['Elastic IP (EIP)', '5_bgp, 200Mbit/s, traffic billing', '1', '$3 + traffic'],
         ['VPC', '172.16.0.0/16', '1', 'Free'],
         ['VPC Subnet', '172.16.1.0/24', '1', 'Free'],
         ['Security Group', 'ICMP, SSH(22), API(4000), Prometheus(9090)', '1', 'Free'],
         ['Total', '—', '—', '~$58-65/mo']])
    p = doc.add_paragraph()
    run = p.add_run('Note: '); run.bold = True
    p.add_run('LLM API call costs are billed separately by each provider.')

    # === Chapter 3: 实施步骤 (Implementation Steps) ===
    doc.add_heading('3. Implementation Steps', level=1)
    doc.add_paragraph('The deployment consists of 4 main steps:')
    for s in ['Preparation — Account setup and IAM delegation',
              'Quick Deploy — One-click via RFS',
              'Getting Started — Configure providers and verify',
              'Quick Uninstall — Clean up resources when done']:
        doc.add_paragraph(s, style='List Number')

    # === Chapter 4: 准备工作 (Preparation) ===
    doc.add_heading('4. Preparation', level=1)
    doc.add_heading('4.1 Account Prerequisites', level=2)
    for c in ['Huawei Cloud account with real-name verification completed',
              'Account not in arrears or frozen',
              'RFS (Resource Formation Service) enabled',
              'MaaS service activated (if using MaaS as provider)']:
        doc.add_paragraph(f'☐ {c}')

    doc.add_heading('4.2 IAM Delegation (for Sub-users)', level=2)
    doc.add_paragraph(
        'If using an IAM sub-user, verify the user is in the admin group. '
        'If not using a sub-user, skip this section.')
    steps = [
        'Log in to Huawei Cloud Console → Identity and Access Management → Delegations',
        'Search for rf_admin_trust',
        'If not found, click "Create Delegation", name: rf_admin_trust, type: "Cloud Service", enter "RFS"',
        'Click "Authorize Now", search and select "Tenant Administrator", scope: "All Resources"',
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f'Step {i}: {s}')

    # === Chapter 5: 快速部署 (Quick Deploy) ===
    doc.add_heading('5. Quick Deployment', level=1)

    doc.add_heading('5.1 Parameter Reference', level=2)
    add_styled_table(doc,
        ['Parameter', 'Required', 'Default', 'Description'],
        [['ecs_password', 'Yes', '—', 'ECS root password, 8-26 chars'],
         ['db_password', 'Yes', '—', 'PostgreSQL password'],
         ['master_key', 'Yes', '—', 'LiteLLM master key, must start with sk-'],
         ['salt_key', 'No', 'auto', 'Encryption key, permanent once set'],
         ['ecs_flavor', 'No', 'x1.4u.8g', 'Flexus X 8vCPU 16GB (enable perf mode in console)'],
         ['system_disk_size', 'No', '200', 'System disk size (GB), high-IO SSD, 40-1024'],
         ['bandwidth_size', 'No', '200', 'EIP bandwidth (Mbit/s), 1-300']])

    doc.add_heading('5.2 Deployment Steps', level=2)
    steps = [
        'Log in to Huawei Cloud RFS Console → Create Resource Stack',
        'Upload template deploying-litellm.tf (or use solution practice URL)',
        'Fill in required parameters (ecs_password, db_password, master_key)',
        'Click "Deploy" and confirm',
        'Wait for stack creation (~1 min) and deploy script (~10 min)',
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f'Step {i}: {s}')

    doc.add_heading('5.3 Deployment Verification', level=2)
    doc.add_paragraph('After ~10 minutes, SSH into the ECS and run:')
    add_code_block(doc, 'docker ps')
    doc.add_paragraph('Expected: 3 containers running (litellm-proxy, litellm-db, litellm-prometheus)')
    add_code_block(doc, 'curl http://localhost:4000/health/liveliness')
    doc.add_paragraph('Expected: {"status": "I\'m alive!"}')
    add_code_block(doc, 'tail -5 /var/log/litellm-bootstrap.log')
    doc.add_paragraph('Check for any error messages.')

    # === Chapter 6: 开始使用 (Getting Started) ===
    doc.add_heading('6. Getting Started', level=1)

    doc.add_heading('6.1 Access Endpoints', level=2)
    add_styled_table(doc,
        ['Service', 'URL', 'Description'],
        [['Admin UI', 'http://<EIP>:4000/ui', 'Login with master_key'],
         ['Chat API', 'http://<EIP>:4000/v1/chat/completions', 'OpenAI-compatible'],
         ['Model List', 'http://<EIP>:4000/v1/models', 'List configured models'],
         ['Health', 'http://<EIP>:4000/health/liveliness', 'Readiness probe'],
         ['Prometheus', 'http://<EIP>:9090', 'Metrics dashboard']])

    doc.add_heading('6.2 Add Provider API Keys', level=2)
    add_code_block(doc,
        'ssh root@<EIP>\n'
        'vi /opt/litellm/.env\n'
        '# Add provider keys:\n'
        '# OPENAI_API_KEY=sk-xxx\n'
        '# ANTHROPIC_API_KEY=sk-ant-xxx\n'
        'cd /opt/litellm && docker compose restart')

    doc.add_heading('6.3 API Call Example', level=2)
    add_code_block(doc,
        'curl http://<EIP>:4000/v1/chat/completions \\\n'
        '  -H "Authorization: Bearer sk-your-master-key" \\\n'
        '  -H "Content-Type: application/json" \\\n'
        '  -d \'{"model": "gpt-4o",\n'
        '       "messages": [{"role": "user", "content": "Hello!"}]}\'')

    doc.add_heading('6.4 Custom Model Configuration', level=2)
    doc.add_paragraph('Edit /opt/litellm/config.yaml:')
    add_code_block(doc,
        'model_list:\n'
        '  - model_name: my-gpt4\n'
        '    litellm_params:\n'
        '      model: openai/gpt-4o\n'
        '      api_key: os.environ/OPENAI_API_KEY\n'
        '  - model_name: my-claude\n'
        '    litellm_params:\n'
        '      model: anthropic/claude-sonnet-4-20250514\n'
        '      api_key: os.environ/ANTHROPIC_API_KEY')
    doc.add_paragraph('After editing: cd /opt/litellm && docker compose restart')

    doc.add_heading('6.5 Virtual Key Management', level=2)
    steps = [
        'Log in to Admin UI at http://<EIP>:4000/ui',
        'Navigate to API Keys section → Create Key',
        'Set key name, budget limits, and allowed models',
        'Distribute generated keys to team members',
        'Monitor spend per key in the Usage dashboard',
    ]
    for s in steps:
        doc.add_paragraph(s, style='List Number')

    doc.add_heading('6.6 Full Chain Verification', level=2)
    add_code_block(doc,
        '# 1. Health check\n'
        'curl -s http://<EIP>:4000/health/liveliness\n'
        '# 2. List models\n'
        'curl -s http://<EIP>:4000/v1/models | head\n'
        '# 3. Send a chat request\n'
        'curl -s http://<EIP>:4000/v1/chat/completions \\\n'
        '  -H "Authorization: Bearer sk-master-key" \\\n'
        '  -H "Content-Type: application/json" \\\n'
        '  -d \'{"model":"gpt-4o","messages":[{"role":"user","content":"hi"}]}\'\n'
        '# 4. Check Prometheus metrics\n'
        'curl -s http://<EIP>:9090/api/v1/query?query=litellm_requests_total')

    # === Chapter 7: 快速卸载 (Quick Uninstall) ===
    doc.add_heading('7. Quick Uninstall', level=1)
    steps = [
        'Log in to RFS Console (https://console.huaweicloud.com/rfs/)',
        'Locate the resource stack corresponding to this deployment',
        'Click "Delete Stack", select "Delete Resources"',
        'Type "Delete" to confirm',
    ]
    for s in steps:
        doc.add_paragraph(s, style='List Number')
    p = doc.add_paragraph()
    run = p.add_run('⚠️ Warning: '); run.bold = True
    p.add_run(
        'Back up /opt/litellm/volumes/db/data before deletion to preserve '
        'virtual keys, spend history, and configuration.')
    add_code_block(doc,
        '# Backup command (run before uninstalling)\n'
        'ssh root@<EIP>\n'
        'tar czf litellm-backup-$(date +%Y%m%d).tar.gz /opt/litellm/volumes/db/data')

    # === Chapter 8: 附录 (Appendix) ===
    doc.add_heading('8. Appendix', level=1)
    doc.add_heading('8.1 Glossary', level=2)
    add_styled_table(doc,
        ['Term', 'Description'],
        [['LiteLLM', 'Open-source LLM API gateway unifying 100+ providers'],
         ['Flexus X', 'Huawei Cloud flexible compute instance for SMBs and developers'],
         ['EIP', 'Elastic IP — independent public IP and bandwidth'],
         ['VPC', 'Virtual Private Cloud — isolated network on Huawei Cloud'],
         ['RFS', 'Resource Formation Service — Terraform-compatible orchestration'],
         ['Prometheus', 'Open-source monitoring system with metrics collection']])

    doc.add_heading('8.2 References', level=2)
    for ref in ['LiteLLM GitHub: https://github.com/BerriAI/litellm',
                'LiteLLM Proxy Docs: https://docs.litellm.ai/docs/proxy',
                'Huawei Cloud RFS: https://support.huaweicloud.com/tr-aos/rf_05_0001.html',
                'Prometheus Docs: https://prometheus.io/docs/']:
        doc.add_paragraph(ref, style='List Bullet')

    # === Chapter 9: 修订记录 (Revision History) ===
    doc.add_heading('9. Revision History', level=1)
    add_styled_table(doc,
        ['Date', 'Version', 'Changes'],
        [['2026-06-10', 'v1.1', 'Upgraded to 8vCPU 16GB Flexus X, 200GB SSD, performance mode, 200Mbit/s bandwidth'],
         ['2026-06-09', 'v1.0', 'Initial release — LiteLLM HK deployment verified (3 containers, API OK)']])

    path = os.path.join(OUT, 'LiteLLM-Deployment-Guide.docx')
    doc.save(path)
    print(f'OK: {path}')


# =============================================================
#  2. 部署指南 (ZH) - Hermes 9-chapter structure
# =============================================================
def gen_deploy_zh():
    doc = Document()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(120)
    run = p.add_run('部署指南'); run.bold = True; run.font.size = Pt(26); run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('LiteLLM，统一的 LLM 网关'); run.font.size = Pt(16); run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

    for line in ['', '', '文档类型：SAC 部署指南', '方案版本：v1.0', '发布日期：2026-06-10', '部署区域：香港 (ap-southeast-1)']:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line); run.font.size = Pt(11); run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_page_break()

    # 1. 方案概述
    doc.add_heading('1. 方案概述', level=1)
    doc.add_heading('1.1 方案简介', level=2)
    doc.add_paragraph(
        '本方案通过单个 Terraform 模板，在华为云 Flexus X 实例上一键部署 LiteLLM（GitHub 17k+ Stars）'
        '——统一 LLM API 网关。LiteLLM 将 100+ 个 LLM 提供商（OpenAI、Anthropic、Azure、Cohere 等）'
        '统一为单个兼容 OpenAI 的 API，内置虚拟密钥管理、花费追踪、负载均衡和管理后台。')

    doc.add_heading('1.2 方案架构', level=2)
    doc.add_paragraph('采用 Docker Compose 编排三个容器：')
    add_styled_table(doc,
        ['容器', '镜像', '作用'],
        [['LiteLLM Proxy', 'ghcr.io/berriai/litellm-database:main-stable', 'API 网关，路由，密钥管理'],
         ['PostgreSQL 16', 'postgres:16', '虚拟密钥、用量数据、配置'],
         ['Prometheus', 'prom/prometheus:v2.53.0', '指标采集与监控']])

    doc.add_heading('1.3 核心优势', level=2)
    for b in ['统一 100+ LLM 提供商，兼容 OpenAI API 格式',
              '虚拟密钥管理，按 Key 设预算上限，用量实时追踪',
              '多提供商故障转移与负载均衡',
              'RFS 一键部署，约 10 分钟即可使用',
              '生产级配置 — 8vCPU 16GB Flexus X + 性能模式 + 200GB SSD']:
        doc.add_paragraph(b, style='List Bullet')

    doc.add_heading('1.4 约束与限制', level=2)
    for c in ['需华为云账号且余额充足', '需开通 RFS 服务',
              'IAM 子用户需创建 rf_admin_trust 委托',
              'LLM API 调用费用由各提供商单独计费']:
        doc.add_paragraph(c, style='List Bullet')

    # 2. 资源和成本规划
    doc.add_heading('2. 资源和成本规划', level=1)
    add_styled_table(doc,
        ['华为云服务', '规格', '数量', '月费 (USD)'],
        [['Flexus X 实例', 'x1.4u.8g (8vCPU 16GB), Ubuntu 24.04, 性能模式', '1', '~$41'],
         ['系统盘', '200GB 高IO SSD', '1', '~$14'],
         ['弹性公网 IP', '5_bgp, 200Mbit/s, 按流量', '1', '$3 + 流量'],
         ['VPC', '172.16.0.0/16', '1', '免费'],
         ['子网', '172.16.1.0/24', '1', '免费'],
         ['安全组', 'ICMP, 22, 4000, 9090', '1', '免费'],
         ['合计', '—', '—', '~$58-65/月']])
    p = doc.add_paragraph()
    run = p.add_run('说明：'); run.bold = True
    p.add_run('LLM API 调用费用由各提供商单独计费，不计入上表。')

    # 3. 实施步骤
    doc.add_heading('3. 实施步骤', level=1)
    doc.add_paragraph('部署全流程包含以下 4 个步骤：')
    for s in ['准备工作 — 账号准备与 IAM 委托', '快速部署 — RFS 一键部署',
              '开始使用 — 配置提供商并验证', '快速卸载 — 释放资源']:
        doc.add_paragraph(s, style='List Number')

    # 4. 准备工作
    doc.add_heading('4. 准备工作', level=1)
    doc.add_heading('4.1 账号前置条件', level=2)
    for c in ['已完成华为云账号注册和实名认证',
              '账户不处于欠费或冻结状态',
              '已开通 RFS（资源编排服务）',
              '如使用 MaaS 作为提供商，需开通 MaaS 服务']:
        doc.add_paragraph(f'☐ {c}')

    doc.add_heading('4.2 IAM 委托（IAM 子用户需要）', level=2)
    doc.add_paragraph('如使用 IAM 子用户，需创建 rf_admin_trust 委托。如使用主账号，可跳过本节。')
    steps = [
        '登录华为云控制台 → 统一身份认证 → 委托',
        '搜索 rf_admin_trust',
        '如不存在，单击"创建委托"，名称 rf_admin_trust，类型"云服务"，输入 RFS',
        '单击"立即授权"，选择 Tenant Administrator，作用范围"所有资源"',
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f'步骤 {i}：{s}')

    # 5. 快速部署
    doc.add_heading('5. 快速部署', level=1)

    doc.add_heading('5.1 参数说明', level=2)
    add_styled_table(doc,
        ['参数', '必填', '默认值', '说明'],
        [['ecs_password', '是', '—', 'ECS root 密码，8-26 位'],
         ['db_password', '是', '—', 'PostgreSQL 数据库密码'],
         ['master_key', '是', '—', 'LiteLLM 管理密钥，必须以 sk- 开头'],
         ['salt_key', '否', '自动', 'API Key 加密密钥，设置后不可更改'],
         ['ecs_flavor', '否', 'x1.4u.8g', 'Flexus X 8vCPU 16GB（建议在控制台开启性能模式）'],
         ['system_disk_size', '否', '200', '系统盘大小 (GB)，高IO SSD，范围 40-1024'],
         ['bandwidth_size', '否', '200', '弹性公网带宽 (Mbit/s)，范围 1-300']])

    doc.add_heading('5.2 一键部署步骤', level=2)
    steps = [
        '登录华为云 RFS 控制台 → 创建资源栈',
        '上传模板 deploying-litellm.tf（或使用解决方案实践链接）',
        '填写必填参数（ecs_password, db_password, master_key）',
        '单击"部署"并确认',
        '等待资源栈创建（约 1 分钟）和部署脚本执行（约 10 分钟）',
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f'步骤 {i}：{s}')

    doc.add_heading('5.3 部署验证', level=2)
    doc.add_paragraph('部署完成后 SSH 登录 ECS，执行以下命令验证：')
    add_code_block(doc, 'docker ps')
    doc.add_paragraph('预期输出：3 个容器运行中（litellm-proxy, litellm-db, litellm-prometheus）')
    add_code_block(doc, 'curl http://localhost:4000/health/liveliness')
    doc.add_paragraph('预期输出：{"status": "I\'m alive!"}')

    # 6. 开始使用
    doc.add_heading('6. 开始使用', level=1)

    doc.add_heading('6.1 访问地址', level=2)
    add_styled_table(doc,
        ['服务', '地址', '说明'],
        [['管理后台', 'http://<EIP>:4000/ui', '用 master_key 登录'],
         ['Chat API', 'http://<EIP>:4000/v1/chat/completions', '兼容 OpenAI 格式'],
         ['模型列表', 'http://<EIP>:4000/v1/models', '查看已配置模型'],
         ['健康检查', 'http://<EIP>:4000/health/liveliness', '就绪探测'],
         ['Prometheus', 'http://<EIP>:9090', '监控面板']])

    doc.add_heading('6.2 添加提供商 API Key', level=2)
    add_code_block(doc,
        'ssh root@<EIP>\n'
        'vi /opt/litellm/.env\n'
        '# 添加提供商密钥：\n'
        '# OPENAI_API_KEY=sk-xxx\n'
        '# ANTHROPIC_API_KEY=sk-ant-xxx\n'
        'cd /opt/litellm && docker compose restart')

    doc.add_heading('6.3 API 调用示例', level=2)
    add_code_block(doc,
        'curl http://<EIP>:4000/v1/chat/completions \\\n'
        '  -H "Authorization: Bearer sk-your-master-key" \\\n'
        '  -H "Content-Type: application/json" \\\n'
        '  -d \'{"model": "gpt-4o",\n'
        '       "messages": [{"role": "user", "content": "你好！"}]}\'')

    doc.add_heading('6.4 虚拟密钥管理', level=2)
    steps = [
        '登录管理后台 http://<EIP>:4000/ui',
        '进入 API Keys 页面 → 创建密钥',
        '设置密钥名称、预算上限和允许的模型',
        '将生成的密钥分发给团队各成员',
        '在 Usage 面板查看每个密钥的用量统计',
    ]
    for s in steps:
        doc.add_paragraph(s, style='List Number')

    doc.add_heading('6.5 全链路验证', level=2)
    add_code_block(doc,
        '# 1. 健康检查\n'
        'curl -s http://<EIP>:4000/health/liveliness\n'
        '# 2. 模型列表\n'
        'curl -s http://<EIP>:4000/v1/models | head\n'
        '# 3. 发送 Chat 请求\n'
        'curl -s http://<EIP>:4000/v1/chat/completions \\\n'
        '  -H "Authorization: Bearer sk-master-key" \\\n'
        '  -H "Content-Type: application/json" \\\n'
        '  -d \'{"model":"gpt-4o","messages":[{"role":"user","content":"hi"}]}\'\n'
        '# 4. Prometheus 指标\n'
        'curl -s http://<EIP>:9090/api/v1/query?query=litellm_requests_total')

    # 7. 快速卸载
    doc.add_heading('7. 快速卸载', level=1)
    steps = [
        '登录 RFS 控制台（https://console.huaweicloud.com/rfs/）',
        '找到本方案部署的资源栈',
        '单击"删除资源栈"，选择"删除资源"',
        '输入 Delete 确认删除',
    ]
    for s in steps:
        doc.add_paragraph(s, style='List Number')
    p = doc.add_paragraph()
    run = p.add_run('⚠️ 警告：'); run.bold = True
    p.add_run('删除前请备份 /opt/litellm/volumes/db/data，以保留虚拟密钥、用量历史等数据。')
    add_code_block(doc,
        '# 卸载前备份命令\n'
        'ssh root@<EIP>\n'
        'tar czf litellm-backup-$(date +%Y%m%d).tar.gz /opt/litellm/volumes/db/data')

    # 8. 附录
    doc.add_heading('8. 附录', level=1)
    doc.add_heading('8.1 名词解释', level=2)
    add_styled_table(doc,
        ['术语', '说明'],
        [['LiteLLM', '开源 LLM API 网关，统一 100+ 提供商'],
         ['Flexus X', '华为云面向中小企业和开发者的柔性算力实例'],
         ['EIP', '弹性公网 IP，提供独立公网地址和带宽'],
         ['VPC', '虚拟私有云，隔离的虚拟网络环境'],
         ['RFS', '资源编排服务，兼容 Terraform 的编排服务'],
         ['Prometheus', '开源监控系统，支持指标采集和查询']])

    doc.add_heading('8.2 参考资源', level=2)
    for ref in ['LiteLLM GitHub：https://github.com/BerriAI/litellm',
                'LiteLLM Proxy 文档：https://docs.litellm.ai/docs/proxy',
                '华为云 RFS：https://support.huaweicloud.com/tr-aos/rf_05_0001.html',
                'Prometheus 文档：https://prometheus.io/docs/']:
        doc.add_paragraph(ref, style='List Bullet')

    # 9. 修订记录
    doc.add_heading('9. 修订记录', level=1)
    add_styled_table(doc,
        ['发布日期', '版本', '修订内容'],
        [['2026-06-10', 'v1.1', '升级至 8vCPU 16GB Flexus X、200GB SSD、性能模式、200Mbit/s 带宽'],
         ['2026-06-09', 'v1.0', '初始版本，LiteLLM HK 部署验证通过（3 容器运行正常，API 验证通过）']])

    path = os.path.join(OUT, 'LiteLLM-部署指南.docx')
    doc.save(path)
    print(f'OK: {path}')


# =============================================================
#  3. Business Value Report (EN)
# =============================================================
def gen_biz_en():
    doc = Document()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(120)
    run = p.add_run('Business Value Report'); run.bold = True; run.font.size = Pt(26); run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('LiteLLM, The Unified LLM Gateway'); run.font.size = Pt(16); run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
    for line in ['', '', 'Document Type: SAC Business Value Report',
                 'Solution Version: v1.0', 'Release Date: 2026-06-10', 'Region: Hong Kong (ap-southeast-1)']:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line); run.font.size = Pt(11); run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_page_break()

    doc.add_heading('1. Product Overview', level=1)
    p = doc.add_paragraph()
    run = p.add_run('One-liner: '); run.bold = True
    p.add_run('A one-click deployed LLM API gateway that unifies 100+ providers under a single OpenAI-compatible endpoint, with virtual key management, spend tracking, and high-availability routing.')
    doc.add_heading('1.1 Enterprise Pain Points', level=2)
    doc.add_heading('Pain 1: Multi-Provider Chaos', level=3)
    add_styled_table(doc, ['Issue', 'Impact'],
        [['API Fragmentation', 'Each provider has different auth, SDK, endpoints — integration cost grows linearly'],
         ['Key Sprawl', 'API keys scattered across code and config files — security risk'],
         ['Vendor Lock-in', 'Switching providers requires code changes, locks you to one vendor']])
    doc.add_heading('Pain 2: Cost Blindness', level=3)
    add_styled_table(doc, ['Issue', 'Impact'],
        [['No Cost Attribution', 'Monthly bills show total only — no per-team breakdown'],
         ['No Budget Control', 'No per-key limits — abnormal traffic causes bill spikes'],
         ['No Chargeback', 'Teams share keys — costs cannot be allocated']])
    doc.add_heading('Pain 3: No High Availability', level=3)
    add_styled_table(doc, ['Issue', 'Impact'],
        [['Single Provider Dependency', 'Provider outage = business outage'],
         ['No Auto Failover', 'No automatic switch to backup providers'],
         ['No Rate Limiting', 'No unified throttling — risk of being rate-limited']])
    doc.add_heading('1.2 Value Proposition', level=2)
    doc.add_paragraph('Unified API Gateway + Virtual Key Management + Multi-Provider Routing + One-Click Deploy = Lower Cost + Higher Security + Stronger Availability + Full Transparency')

    doc.add_heading('2. Business Scenarios', level=1)
    doc.add_heading('Scenario 1: Multi-Team LLM Management', level=2)
    doc.add_paragraph('Target: Mid-to-large enterprises (50-500 staff), multiple teams using LLMs.')
    doc.add_paragraph('Through LiteLLM\'s virtual key mechanism, create independent keys per team/project, set budget caps, and track usage automatically.')
    add_styled_table(doc, ['Metric', 'Traditional', 'With LiteLLM', 'Improvement'],
        [['API Key Management', 'Scattered in code', 'Centralized console', 'Centralized'],
         ['Key Rotation Cost', 'Code change per key', 'One-click in console', '90% labor saved'],
         ['Cost Attribution', 'Impossible', 'Per-key, per-request', 'Full transparency'],
         ['Abuse Prevention', 'None', 'Auto-block on budget', 'Zero-trust security']])

    doc.add_heading('Scenario 2: Provider Switching & Cost Optimization', level=2)
    doc.add_paragraph('Target: Cost-conscious enterprises optimizing LLM model selection.')
    add_styled_table(doc, ['Model', 'Input ($/M tokens)', 'Output ($/M tokens)', 'Best For'],
        [['GPT-4o-mini', '$0.15', '$0.60', 'Classification, summarization'],
         ['DeepSeek-V3', '$0.27', '$1.10', 'Code generation, analysis'],
         ['GPT-4o', '$2.50', '$10.00', 'Complex reasoning'],
         ['Claude Sonnet 4', '$3.00', '$15.00', 'Architecture design']])
    p = doc.add_paragraph(); run = p.add_run('Business Impact: '); run.bold = True; p.add_run('30-50% reduction in annual LLM costs.')

    doc.add_heading('Scenario 3: High-Availability LLM Infrastructure', level=2)
    doc.add_paragraph('Target: Enterprises integrating AI into core business processes requiring SLA guarantees.')
    add_styled_table(doc, ['Metric', 'Single Provider', 'LiteLLM Multi-Provider'],
        [['Availability', '99%', '99.9%+'], ['Annual Downtime', '~87 hours', '~8 hours'],
         ['Failover', 'Manual, hours', 'Automatic, seconds']])

    doc.add_heading('Scenario 4: Dev/Test Multi-Model Compatibility', level=2)
    doc.add_paragraph('Target: AI application dev teams using different models across dev/test/prod.')
    doc.add_paragraph('50%+ dev efficiency gain — one codebase for all environments, zero code changes to switch.')

    doc.add_heading('3. Core Advantages', level=1)
    for t, d in [('100+ LLM Providers, One API', 'Single OpenAI-compatible endpoint. Switch providers via YAML config. No code changes.'),
                  ('Virtual Key Management', 'Per-team/per-app keys with budget caps, auto-blocking, and full usage tracking. 80% security risk reduction.'),
                  ('Smart Routing & HA', 'Latency-based routing, automatic failover, load balancing, rate limiting.'),
                  ('One-Click Deploy', 'Single Terraform template, 3 auto-configured containers, ~10 min deployment.'),
                  ('Production-Ready', '8vCPU 16GB Flexus X with performance mode. ~$58-65/mo (Hong Kong). 200GB SSD.'),
                  ('Built-in Monitoring', 'Prometheus metrics + Admin UI dashboard. Usage charts, cost trends, model heatmaps.')]:
        doc.add_heading(f'Advantage: {t}', level=2); doc.add_paragraph(d)

    doc.add_heading('4. Customer Benefits Summary', level=1)
    add_styled_table(doc, ['Dimension', 'Metric', 'Business Value'],
        [['Cost Optimization', '30-50% reduction', 'Save $10K-$100K+ annually'],
         ['Security', '80% risk reduction', 'Centralized keys + virtual isolation'],
         ['Dev Efficiency', '50%+ faster integration', 'One API for all providers'],
         ['Availability', '99% → 99.9%+', '< 9 hrs downtime per year'],
         ['Deployment Speed', '~10 min one-click', 'Decision to go-live in hours'],
         ['Cost Transparency', 'Per-key, per-request', 'Clear chargeback, no black-box bills']])

    path = os.path.join(OUT, 'LiteLLM-Business-Value-Report.docx'); doc.save(path); print(f'OK: {path}')


# =============================================================
#  4. Business Value Report (ZH)
# =============================================================
def gen_biz_zh():
    doc = Document()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(120)
    run = p.add_run('业务价值报告'); run.bold = True; run.font.size = Pt(26); run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('LiteLLM，统一的 LLM 网关'); run.font.size = Pt(16); run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
    for line in ['', '', '文档类型：SAC 业务价值报告', '方案版本：v1.0', '发布日期：2026-06-10', '部署区域：香港 (ap-southeast-1)']:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line); run.font.size = Pt(11); run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_page_break()

    doc.add_heading('1. 产品概述', level=1)
    p = doc.add_paragraph(); run = p.add_run('一句话定位：'); run.bold = True
    p.add_run('一键部署的 LLM API 统一网关，通过虚拟密钥管理和多提供商路由，让企业以单个兼容 OpenAI 的端点管理所有大模型调用，实现成本透明、密钥可控、服务高可用。')

    doc.add_heading('1.1 企业痛点', level=2)
    doc.add_heading('痛点一：多提供商管理混乱', level=3)
    add_styled_table(doc, ['问题', '影响'],
        [['API 碎片化', '每个提供商独立 API、认证、SDK，集成成本线性增长'],
         ['密钥分散', '密钥散落在代码和配置中，安全隐患大'],
         ['厂商锁定', '切换提供商需改代码，被单一绑定']])
    doc.add_heading('痛点二：成本失控', level=3)
    add_styled_table(doc, ['问题', '影响'],
        [['费用不透明', '总账单无法拆分到团队/项目'], ['无预算控制', '无按 Key 限额，异常流量导致账单飙升'],
         ['无法分摊', '多团队共享 Key，费用归属不清']])
    doc.add_heading('痛点三：缺乏高可用', level=3)
    add_styled_table(doc, ['问题', '影响'],
        [['单点依赖', '提供商故障 = 业务中断'], ['无故障转移', '主提供商不可用时不会自动切换'],
         ['无限速', '无统一限流，突发流量被封']])
    doc.add_heading('1.2 价值主张', level=2)
    doc.add_paragraph('统一 API 网关 + 虚拟密钥管理 + 多提供商路由 + 一键部署 = 更低成本 + 更高安全 + 更强可用 + 完全透明')

    doc.add_heading('2. 应用场景（业务视角）', level=1)
    doc.add_heading('场景一：多业务线统一 LLM 管理', level=2)
    doc.add_paragraph('适合客户：中大型企业（50-500 人），多个业务线同时使用 LLM')
    doc.add_paragraph('通过 LiteLLM 虚拟密钥机制，为每个团队/项目创建独立 Key，设置预算上限，自动追踪用量。')
    add_styled_table(doc, ['指标', '传统方式', '本方案', '改善'],
        [['Key 管理', '分散在代码中', '统一后台', '集中化'], ['密钥轮换', '每 Key 改代码部署', '后台一键轮换', '节省 90% 人力'],
         ['成本归属', '无法追踪', '精确到每 Key', '完全透明'], ['违规拦截', '无', '超预算自动拦截', '零信任安全']])
    p = doc.add_paragraph(); run = p.add_run('业务收益：'); run.bold = True; p.add_run('安全风险降低 80%，成本可追溯至每个团队。')

    doc.add_heading('场景二：多提供商灵活切换与成本优化', level=2)
    doc.add_paragraph('适合客户：对 LLM 成本敏感、希望优化模型选型的企业')
    add_styled_table(doc, ['模型', '输入价 ($/M Token)', '输出价 ($/M Token)', '适用场景'],
        [['GPT-4o-mini', '$0.15', '$0.60', '分类、摘要'],
         ['DeepSeek-V3', '$0.27', '$1.10', '代码生成、数据分析'],
         ['GPT-4o', '$2.50', '$10.00', '复杂推理'],
         ['Claude Sonnet 4', '$3.00', '$15.00', '架构设计']])
    p = doc.add_paragraph(); run = p.add_run('业务收益：'); run.bold = True; p.add_run('年 LLM 费用降低 30-50%。')

    doc.add_heading('场景三：高可用 LLM 基础设施', level=2)
    doc.add_paragraph('适合客户：将 AI 集成到核心业务流程、需要 SLA 保障的企业')
    add_styled_table(doc, ['指标', '单一提供商', 'LiteLLM 多提供商'],
        [['可用性', '99%', '99.9%+'], ['年停机', '~87 小时', '~8 小时'], ['故障切换', '手动，数小时', '自动，秒级']])

    doc.add_heading('场景四：开发测试多模型兼容', level=2)
    doc.add_paragraph('适合客户：AI 应用开发团队，需在不同环境使用不同模型')
    doc.add_paragraph('一套代码适用所有环境，切换模型零代码改动。开发效率提升 50%+。')

    doc.add_heading('3. 核心优势', level=1)
    for t, d in [('100+ 提供商统一接入', '单个 OpenAI 兼容端点。切换提供商只需改 YAML，无需改代码。'),
                  ('虚拟密钥管理', '每团队独立 Key，设预算上限，超限自动拦截。风险降低 80%。'),
                  ('智能路由与高可用', '延迟优先路由、自动故障转移、负载均衡、速率限制四层保障。'),
                  ('一键部署', '单个 Terraform 模板，3 容器自动编排，10 分钟部署。'),
                  ('生产级配置', '8vCPU 16GB Flexus X + 性能模式，200GB SSD，月费约 $58-65（香港）。'),
                  ('内置监控', 'Prometheus + 管理后台，用量图表、费用趋势一目了然。')]:
        doc.add_heading(f'优势：{t}', level=2); doc.add_paragraph(d)

    doc.add_heading('4. 客户收益总览', level=1)
    add_styled_table(doc, ['维度', '量化指标', '业务价值'],
        [['成本优化', '降低 30-50%', '年省数万至数十万元'], ['安全提升', '风险降低 80%', '集中密钥 + 虚拟隔离'],
         ['开发效率', '提升 50%+', '一套 API 兼容所有提供商'], ['可用性保障', '99% → 99.9%+', '年停机 < 9 小时'],
         ['部署提效', '一键 10 分钟', '决策到上线从天到小时'], ['管理透明', '每 Key 每请求追踪', '消灭黑盒账单']])

    path = os.path.join(OUT, 'LiteLLM-业务价值报告.docx'); doc.save(path); print(f'OK: {path}')


if __name__ == '__main__':
    gen_deploy_en()
    gen_deploy_zh()
    gen_biz_en()
    gen_biz_zh()
    print('\nAll 4 DOCX files generated!')
    print('  Structure: Deploy Guide follows Hermes 9-chapter pattern')
    print('  Renamed: 技术交付报告 → 部署指南')
