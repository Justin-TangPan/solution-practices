#!/usr/bin/env python3
"""Generate LiteLLM HK DOCX deployment guides (Chinese + English)"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def create_docx(output_path, lang):
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

    # Normal style
    style = doc.styles['Normal']
    font = style.font
    if lang == 'zh':
        font.name = 'Microsoft YaHei'
        style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    else:
        font.name = 'Calibri'
    font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)

    def add_heading(text, level=1):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        return h

    def add_para(text, bold=False, align=None):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        if align:
            p.alignment = align
        return p

    def add_code(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        return p

    # === Title Page ===
    for _ in range(6):
        doc.add_paragraph()

    if lang == 'zh':
        add_para('Litellm — 统一 LLM 网关',
                 bold=True, align=WD_ALIGN_PARAGRAPH.CENTER).runs[0].font.size = Pt(28)
        add_para('中国香港 一键部署指南',
                 bold=True, align=WD_ALIGN_PARAGRAPH.CENTER).runs[0].font.size = Pt(22)
        doc.add_paragraph()
        add_para('Solution Practices (SAC)',
                 align=WD_ALIGN_PARAGRAPH.CENTER).runs[0].font.size = Pt(14)
        add_para('华为云资源编排服务 (RFS)',
                 align=WD_ALIGN_PARAGRAPH.CENTER).runs[0].font.size = Pt(14)
    else:
        add_para('LiteLLM — Unified LLM Gateway',
                 bold=True, align=WD_ALIGN_PARAGRAPH.CENTER).runs[0].font.size = Pt(28)
        add_para('Hong Kong Region — One-Click Deployment Guide',
                 bold=True, align=WD_ALIGN_PARAGRAPH.CENTER).runs[0].font.size = Pt(22)
        doc.add_paragraph()
        add_para('Solution Practices (SAC)',
                 align=WD_ALIGN_PARAGRAPH.CENTER).runs[0].font.size = Pt(14)
        add_para('Huawei Cloud Resource Formation Service (RFS)',
                 align=WD_ALIGN_PARAGRAPH.CENTER).runs[0].font.size = Pt(14)

    doc.add_page_break()

    # === TOC ===
    if lang == 'zh':
        add_para('目录', bold=True,
                 align=WD_ALIGN_PARAGRAPH.CENTER).runs[0].font.size = Pt(16)
        doc.add_paragraph()
        toc_items = [
            '1. 解决方案概述',
            '2. 架构',
            '3. 适用场景与核心优势',
            '4. 部署指南',
            '5. 参数说明',
            '6. 快速使用',
            '7. 预估费用',
            '8. 快速卸载',
            '9. 更多资源',
        ]
    else:
        add_para('Table of Contents', bold=True,
                 align=WD_ALIGN_PARAGRAPH.CENTER).runs[0].font.size = Pt(16)
        doc.add_paragraph()
        toc_items = [
            '1. Solution Overview',
            '2. Architecture',
            '3. Use Cases & Key Benefits',
            '4. Deployment Guide',
            '5. Parameters',
            '6. Getting Started',
            '7. Estimated Cost',
            '8. Quick Uninstall',
            '9. More Resources',
        ]
    for item in toc_items:
        add_para(item)

    doc.add_page_break()

    # === Section 1: Solution Overview ===
    if lang == 'zh':
        add_heading('1. 解决方案概述', 1)
        add_para(
            'LiteLLM 是一个开源的统一 LLM API '
            '网关（GitHub 17k+ Stars），通过单个 '
            'Terraform 模板在华为云 Flexus X 实例上'
            '一键部署。LiteLLM 将 100+ 个 LLM 提供商'
            '（OpenAI、Anthropic、Azure、Cohere 等）统一为'
            '单个兼容 OpenAI 格式的 API，内置'
            '虚拟密钥管理、花费追踪、'
            '负载均衡和管理后台。')
        add_para(
            '本方案面向中国香港区域'
            '（ap-southeast-1）部署，使用 Docker Hub '
            '官方镜像直连，无需华为云 SWR '
            '镜像仓库。')
    else:
        add_heading('1. Solution Overview', 1)
        add_para(
            'LiteLLM is an open-source unified LLM API gateway (17k+ Stars on GitHub) '
            'deployed on a Huawei Cloud Flexus X instance via a single Terraform template. '
            'It unifies 100+ LLM providers (OpenAI, Anthropic, Azure, Cohere, etc.) into '
            'a single OpenAI-compatible API with built-in virtual key management, spend '
            'tracking, load balancing, and an Admin Web UI.')
        add_para(
            'This solution is designed for the Hong Kong region (ap-southeast-1), '
            'using direct Docker Hub image pulls without Huawei Cloud SWR.')

    # === Section 2: Architecture ===
    if lang == 'zh':
        add_heading('2. 架构', 1)
        add_para('部署资源概览：')
    else:
        add_heading('2. Architecture', 1)
        add_para('Deployed Resources Overview:')

    table = doc.add_table(rows=6, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Light Grid Accent 1'
    headers = ['资源', '规格', '数量', '说明'] \
        if lang == 'zh' else ['Resource', 'Spec', 'Qty', 'Description']

    if lang == 'zh':
        data = [
            ['Flexus X 实例',
             'x1.2u.4g (2vCPUs 4GiB)', '1',
             '运行 LiteLLM + PostgreSQL + Prometheus'],
            ['弹性公网 IP (EIP)',
             '5-bgp, 300Mbit/s 按流量', '1',
             'API 网关入口'],
            ['VPC', '172.16.0.0/16', '1', '网络隔离'],
            ['VPC 子网', '172.16.1.0/24', '1',
             '内部通信'],
            ['安全组', '-', '1',
             '开放 HTTP(4000) + Prometheus(9090) + SSH(22)'],
        ]
    else:
        data = [
            ['Flexus X Instance',
             'x1.2u.4g (2vCPUs 4GiB)', '1',
             'Runs LiteLLM + PostgreSQL + Prometheus'],
            ['Elastic IP (EIP)',
             '5-bgp, 300Mbit/s traffic', '1',
             'API gateway entry'],
            ['VPC', '172.16.0.0/16', '1', 'Network isolation'],
            ['VPC Subnet', '172.16.1.0/24', '1',
             'Internal communication'],
            ['Security Group', '-', '1',
             'Opens HTTP(4000) + Prometheus(9090) + SSH(22)'],
        ]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row_data in enumerate(data):
        for ci, val in enumerate(row_data):
            table.rows[ri + 1].cells[ci].text = val

    add_para('')
    if lang == 'zh':
        add_para(
            'LiteLLM 提供兼容 OpenAI 格式的 API，'
            '底层可路由到任意 LLM 提供商。')
    else:
        add_para(
            'LiteLLM provides an OpenAI-compatible API that routes to '
            'any LLM provider.')

    # === Section 3: Use Cases & Benefits ===
    if lang == 'zh':
        add_heading('3. 适用场景与核心优势', 1)
        add_heading('适用场景', 2)
        scenarios = [
            ('需要统一访问多个 LLM 提供商'
             '并集中管理密钥的团队'),
            ('需要虚拟密钥管理、花费追踪'
             '和成本分配的组织'),
            ('需要负载均衡、故障转移和'
             '速率限制的 LLM 应用'),
            ('使用 OpenAI 兼容 API 但需要在多个'
             '提供商间切换的项目'),
            ('AI 应用的开发和测试环境'),
        ]
        for s in scenarios:
            doc.add_paragraph(s, style='List Bullet')
        add_heading('核心优势', 2)
        benefits = [
            ('17k+ Stars 开源项目',
             'MIT 许可证，活跃社区，'
             '支持 100+ LLM 提供商'),
            ('单个 Terraform 模板',
             '一个 .tf 文件部署全部组件'
             '，无外部依赖'),
            ('Docker Compose 编排',
             '3 个容器自动配置，'
             '约 10 分钟完成部署'),
            ('兼容 OpenAI 格式的 API',
             '统一的 /v1/chat/completions 端点，'
             '切换提供商无需修改代码'),
            ('管理后台',
             'Web UI 可视化配置模型、'
             'API Key 和用量统计'),
            ('轻量级',
             '仅需 2vCPU 4GiB 即可运行全部服务'),
        ]
    else:
        add_heading('3. Use Cases & Key Benefits', 1)
        add_heading('Use Cases', 2)
        scenarios = [
            'Teams needing unified access to multiple LLM providers with centralized key management',  # noqa
            'Organizations requiring virtual key management, spend tracking, and cost allocation',  # noqa
            'LLM applications needing load balancing, failover, and rate limiting',
            'Projects using OpenAI-compatible API but switching between multiple providers',
            'Development and testing environments for AI applications',
        ]
        for s in scenarios:
            doc.add_paragraph(s, style='List Bullet')
        add_heading('Key Benefits', 2)
        benefits = [
            ('17k+ Stars Open Source',
             'MIT license, active community, supports 100+ LLM providers'),
            ('Single Terraform Template',
             'One .tf file deploys everything — no external dependencies'),
            ('Docker Compose Orchestration',
             '3 containers auto-configured, ~10 min deployment'),
            ('OpenAI-Compatible API',
             'Unified /v1/chat/completions endpoint, zero code changes to switch'),
            ('Admin Web UI',
             'Visual management of models, API keys, and usage statistics'),
            ('Lightweight',
             'Only requires 2vCPU 4GiB to run all services'),
        ]

    table2 = doc.add_table(rows=len(benefits) + 1, cols=2)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    table2.style = 'Light Grid Accent 1'
    bh = ['优势', '说明'] \
        if lang == 'zh' else ['Benefit', 'Description']
    for i, h in enumerate(bh):
        cell = table2.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, (title, desc) in enumerate(benefits):
        table2.rows[ri + 1].cells[0].text = title
        table2.rows[ri + 1].cells[1].text = desc

    # === Section 4: Deployment Guide ===
    if lang == 'zh':
        add_heading('4. 部署指南', 1)
        add_heading('前提条件', 2)
        prereqs = [
            '华为云账号，余额充足',
            '已开通 RFS（资源编排服务）',
            '推荐 Flexus X 实例规格 x1.2u.4g 及以上',
        ]
        for p in prereqs:
            doc.add_paragraph(p, style='List Bullet')
        add_heading('一键部署步骤', 2)
        steps = [
            '登录华为云 RFS 控制台 → '
            '创建资源栈',
            '上传模板 deploying-litellm.tf'
            '（或使用 RFS 链接直接跳转）',
            '配置部署参数（master_key 必须以 '
            'sk- 开头）',
            '点击“部署”',
            '等待部署完成（约 10 分钟，'
            '主要为 Docker 镜像拉取时间）',
        ]
        for i, step in enumerate(steps, 1):
            add_para(f'步骤 {i}：{step}')
    else:
        add_heading('4. Deployment Guide', 1)
        add_heading('Prerequisites', 2)
        prereqs = [
            'Huawei Cloud account with sufficient balance',
            'RFS (Resource Formation Service) enabled',
            'Flexus X instance x1.2u.4g or above recommended',
        ]
        for p in prereqs:
            doc.add_paragraph(p, style='List Bullet')
        add_heading('One-Click Deploy Steps', 2)
        steps = [
            'Log in to Huawei Cloud RFS Console → Create Resource Stack',
            'Upload the deploying-litellm.tf template (or use direct RFS link)',
            'Configure parameters (master_key must start with sk-)',
            'Click "Deploy"',
            'Wait for completion (~10 minutes, mainly Docker image pulls)',
        ]
        for i, step in enumerate(steps, 1):
            add_para(f'Step {i}: {step}')

    # === Section 5: Parameters ===
    add_heading('5. 参数说明' if lang == 'zh' else '5. Parameters', 1)

    if lang == 'zh':
        params = [
            ['solution_name', 'litellm-llm-gateway',
             '解决方案名称，用于所有'
             '资源命名'],
            ['ecs_flavor', 'x1.2u.4g',
             '云服务器规格，推荐 2vCPUs 4GiB '
             '及以上'],
            ['ecs_password', '(必填)',
             '云服务器 root 密码'],
            ['db_password', '(必填)',
             'PostgreSQL 数据库密码'],
            ['master_key', '(必填)',
             'LiteLLM 管理密钥，须以 sk- 开头'],
            ['salt_key', '(自动生成)',
             'API Key 加密密钥，一旦使用不可'
             '更改'],
            ['system_disk_size', '40',
             '系统盘大小（GB），范围 40-1024'],
            ['bandwidth_size', '300',
             '弹性公网带宽（Mbit/s），范围 '
             '1-300'],
        ]
    else:
        params = [
            ['solution_name', 'litellm-llm-gateway',
             'Solution name for all resource naming'],
            ['ecs_flavor', 'x1.2u.4g',
             'ECS flavor, 2vCPUs 4GiB or above recommended'],
            ['ecs_password', '(required)', 'ECS root password'],
            ['db_password', '(required)', 'PostgreSQL password'],
            ['master_key', '(required)',
             'LiteLLM master key, must start with sk-'],
            ['salt_key', '(auto)',
             'Encryption key, CANNOT be changed once used'],
            ['system_disk_size', '40',
             'System disk size (GB), range 40-1024'],
            ['bandwidth_size', '300',
             'EIP bandwidth (Mbit/s), range 1-300'],
        ]

    pheaders = ['参数', '默认值', '说明'] \
        if lang == 'zh' else ['Parameter', 'Default', 'Description']
    table3 = doc.add_table(rows=len(params) + 1, cols=3)
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER
    table3.style = 'Light Grid Accent 1'
    for i, h in enumerate(pheaders):
        cell = table3.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row_data in enumerate(params):
        for ci, val in enumerate(row_data):
            table3.rows[ri + 1].cells[ci].text = val

    # === Section 6: Getting Started ===
    if lang == 'zh':
        add_heading('6. 快速使用', 1)
        add_para('部署完成后，通过浏览器'
                 '访问以下地址：')
        add_code(
            '管理后台:  http://<EIP>:4000/ui')
        add_code(
            'API:       http://<EIP>:4000/v1/chat/completions')
        add_code(
            '模型列表:  http://<EIP>:4000/v1/models')
        add_code(
            '健康检查:  http://<EIP>:4000/health/liveliness')
        add_code(
            '监控:      http://<EIP>:9090')
        add_para('')
        add_heading('添加 API Key', 2)
        add_para('部署后，至少添加一个 LLM '
                 '提供商的 API Key：')
        add_code('ssh root@<EIP>')
        add_code('vi /opt/litellm/.env')
        add_code('# 添加所需的 API Key：')
        add_code('# OPENAI_API_KEY=sk-xxx')
        add_code('# ANTHROPIC_API_KEY=sk-ant-xxx')
        add_code('cd /opt/litellm && docker compose restart')
        add_para('')
        add_heading('API 调用示例', 2)
        add_code(
            'curl http://<EIP>:4000/v1/chat/completions \\')
        add_code(
            '  -H "Authorization: Bearer sk-your-master-key" \\')
        add_code(
            '  -H "Content-Type: application/json" \\')
        add_code(
            '  -d \'{"model": "gpt-4o", "messages": '
            '[{"role": "user", "content": "你好！"}]}\'')
        add_para('')
        add_heading('自定义模型配置', 2)
        add_para('编辑 /opt/litellm/config.yaml 添加自定义'
                 '模型路由：')
        add_code('model_list:')
        add_code('  - model_name: my-gpt4')
        add_code('    litellm_params:')
        add_code('      model: openai/gpt-4o')
        add_code('      api_key: os.environ/OPENAI_API_KEY')
    else:
        add_heading('6. Getting Started', 1)
        add_para('After deployment, access via browser:')
        add_code('Admin UI:   http://<EIP>:4000/ui')
        add_code('API:        http://<EIP>:4000/v1/chat/completions')
        add_code('Models:     http://<EIP>:4000/v1/models')
        add_code('Health:     http://<EIP>:4000/health/liveliness')
        add_code('Prometheus: http://<EIP>:9090')
        add_para('')
        add_heading('Add Provider API Key', 2)
        add_para('After deployment, add at least one LLM provider API key:')
        add_code('ssh root@<EIP>')
        add_code('vi /opt/litellm/.env')
        add_code('# Add the API keys you need:')
        add_code('# OPENAI_API_KEY=sk-xxx')
        add_code('# ANTHROPIC_API_KEY=sk-ant-xxx')
        add_code('cd /opt/litellm && docker compose restart')
        add_para('')
        add_heading('API Call Example', 2)
        add_code(
            'curl http://<EIP>:4000/v1/chat/completions \\')
        add_code(
            '  -H "Authorization: Bearer sk-your-master-key" \\')
        add_code(
            '  -H "Content-Type: application/json" \\')
        add_code(
            '  -d \'{"model": "gpt-4o", "messages": '
            '[{"role": "user", "content": "Hello!"}]}\'')
        add_para('')
        add_heading('Custom Model Configuration', 2)
        add_para(
            'Edit /opt/litellm/config.yaml to add custom model routing:')
        add_code('model_list:')
        add_code('  - model_name: my-gpt4')
        add_code('    litellm_params:')
        add_code('      model: openai/gpt-4o')
        add_code('      api_key: os.environ/OPENAI_API_KEY')

    # === Section 7: Estimated Cost ===
    if lang == 'zh':
        add_heading('7. 预估费用', 1)
        add_para('以下为按需计费模式下的'
                 '预估费用（中国香港区域）：')
        cost_data = [
            ['Flexus X 实例', 'x1.2u.4g',
             '~$0.04-0.07/小时', '~$22-30/月'],
            ['弹性公网 IP', '300Mbit/s 按流量',
             '~$0.01/小时（流量另计）',
             '~$3/月（不含流量）'],
            ['系统盘', '40GB SAS', '-', '~$3/月'],
            ['合计', '', '~$0.05-0.08/小时', '~$28-36/月'],
        ]
    else:
        add_heading('7. Estimated Cost', 1)
        add_para(
            'Estimated costs for pay-per-use billing (Hong Kong region):')
        cost_data = [
            ['Flexus X Instance', 'x1.2u.4g',
             '~$0.04-0.07/hr', '~$22-30/mo'],
            ['Elastic IP', '300Mbit/s traffic',
             '~$0.01/hr (excl. traffic)', '~$3/mo (excl. traffic)'],
            ['System Disk', '40GB SAS', '-', '~$3/mo'],
            ['Total', '', '~$0.05-0.08/hr', '~$28-36/mo'],
        ]

    cheaders = ['资源', '规格', '按需', '包月'] \
        if lang == 'zh' else ['Resource', 'Spec', 'Pay-per-use', 'Monthly']
    table4 = doc.add_table(rows=len(cost_data) + 1, cols=4)
    table4.alignment = WD_TABLE_ALIGNMENT.CENTER
    table4.style = 'Light Grid Accent 1'
    for i, h in enumerate(cheaders):
        cell = table4.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row_data in enumerate(cost_data):
        for ci, val in enumerate(row_data):
            cell = table4.rows[ri + 1].cells[ci]
            cell.text = val
            if ri == len(cost_data) - 1:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True

    if lang == 'zh':
        add_para('')
        add_para(
            '注意：LLM API 调用费用由各提供'
            '商单独计费，不在上述费用'
            '范围内。')

    # === Section 8: Quick Uninstall ===
    if lang == 'zh':
        add_heading('8. 快速卸载', 1)
        uninstall_steps = [
            '登录华为云 RFS 控制台',
            '找到对应资源栈',
            '点击“删除资源栈” → '
            '输入 "Delete" → 确认删除',
        ]
        for i, step in enumerate(uninstall_steps, 1):
            add_para(f'{i}. {step}')
        add_para('')
        add_para(
            '警告：删除前请备份 '
            '/opt/litellm/volumes/db/data 以保留虚拟密钥'
            '和用量数据。')
    else:
        add_heading('8. Quick Uninstall', 1)
        uninstall_steps = [
            'Log in to Huawei Cloud RFS Console',
            'Locate the corresponding resource stack',
            'Click "Delete Resource Stack" → Type "Delete" → Confirm',
        ]
        for i, step in enumerate(uninstall_steps, 1):
            add_para(f'{i}. {step}')
        add_para('')
        add_para(
            'Warning: Backup /opt/litellm/volumes/db/data before deletion '
            'to preserve virtual keys and usage data.')

    # === Section 9: More Resources ===
    add_heading('9. 更多资源' if lang == 'zh'
                else '9. More Resources', 1)
    resources = [
        ('LiteLLM GitHub', 'https://github.com/BerriAI/litellm'),
        ('LiteLLM Proxy Docs' if lang == 'en' else 'LiteLLM Proxy 文档',
         'https://docs.litellm.ai/docs/proxy'),
        ('Huawei Cloud RFS' if lang == 'en' else '华为云 RFS',
         'https://support.huaweicloud.com/rfs/'),
    ]
    for name, url in resources:
        add_para(f'{name}: {url}')

    # Footer
    doc.add_paragraph()
    end_text = '— 文档结束 —' if lang == 'zh' \
        else '— End of Document —'
    add_para(end_text, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.save(output_path)
    print(f'OK: {output_path}')


if __name__ == '__main__':
    base = (
        r'C:\Users\Administrator\Desktop\Project\claudeproject'
        r'\solution-practices\release\litellm\hk'
    )
    create_docx(f'{base}\\LiteLLM-部署指南.docx', 'zh')
    create_docx(f'{base}\\LiteLLM-Deployment-Guide.docx', 'en')
    print('Done!')
