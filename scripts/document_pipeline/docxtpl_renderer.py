"""Render a structured document model through a real docxtpl Word master."""
from __future__ import annotations

from copy import deepcopy
import json
from pathlib import Path
import re
from tempfile import TemporaryDirectory
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docxtpl import DocxTemplate
from lxml import etree

from .docx_renderer import W, _content_elements, _paragraph
from .markdown_parser import parse_markdown
from .models import DocumentModel, Section


def _style(document: Document, style_id: str):
    return next(style for style in document.styles if style.style_id == style_id)


def _replace_paragraph(paragraph, text: str) -> None:
    properties = paragraph._p.pPr
    run_properties = deepcopy(paragraph.runs[0]._r.rPr) if paragraph.runs and paragraph.runs[0]._r.rPr is not None else None
    for child in list(paragraph._p):
        if child is not properties:
            paragraph._p.remove(child)
    run = paragraph.add_run(text)
    if run_properties is not None:
        run._r.insert(0, run_properties)


def build_company_template(source: str | Path, output: str | Path) -> Path:
    """Create a visually identical Jinja master while removing sample solution content."""
    source, output = Path(source), Path(output)
    output.parent.mkdir(parents=True, exist_ok=True)
    document = Document(source)
    body = document._body._element
    first_heading = next(
        paragraph._p
        for paragraph in document.paragraphs
        if paragraph.style
        and paragraph.style.style_id == "Heading1"
        and paragraph.text.strip().startswith("方案概述")
    )
    children = list(body)
    final_section = children[-1] if children[-1].tag == qn("w:sectPr") else None
    for element in children[children.index(first_heading) :]:
        if element is not final_section:
            body.remove(element)

    cover_done = False
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if "LiteLLM" in text and not cover_done:
            _replace_paragraph(paragraph, "{{ document_title }}")
            cover_done = True
        elif text.startswith("文档版本"):
            _replace_paragraph(paragraph, "文档版本\t{{ version }}")
        elif text.startswith("发布日期"):
            _replace_paragraph(paragraph, "发布日期\t{{ date }}")

    for section in document.sections:
        for node in section.header._element.iter(qn("w:t")):
            if node.text and "LiteLLM" in node.text:
                node.text = node.text.replace("LiteLLM，统一的 AI 管理网关", "{{ short_title }}")

    for table in document._element.iter(qn("w:tbl")):
        if table.find(qn("w:tblGrid")) is not None:
            continue
        grid = OxmlElement("w:tblGrid")
        rows = table.findall(qn("w:tr"))
        cells = rows[0].findall(qn("w:tc")) if rows else []
        for cell in cells or [None]:
            column = OxmlElement("w:gridCol")
            cell_width = cell.find(f"./{qn('w:tcPr')}/{qn('w:tcW')}") if cell is not None else None
            width = cell_width.get(qn("w:w")) if cell_width is not None else "0"
            column.set(qn("w:w"), width)
            grid.append(column)
        table.insert(1, grid)

    placeholder = document.add_paragraph("{{p body }}", style=_style(document, "BodyText"))._p
    if final_section is not None:
        body.remove(placeholder)
        body.insert(len(body) - 1, placeholder)
    update_fields = document.settings._element.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = parse_xml('<w:updateFields xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="true"/>')
        document.settings._element.append(update_fields)
    document.save(output)
    _add_header_placeholders(output)
    _fix_table_grids(output)
    return output


def _rewrite_members(path: Path, replacements: dict[str, bytes]) -> None:
    with ZipFile(path) as source:
        members = [(info, replacements.get(info.filename, source.read(info.filename))) for info in source.infolist()]
    with ZipFile(path, "w", ZIP_DEFLATED) as target:
        for info, payload in members:
            target.writestr(info, payload)


def _add_header_placeholders(path: Path) -> None:
    namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    with ZipFile(path) as source:
        replacements = {}
        for name in source.namelist():
            if not name.startswith("word/header") or not name.endswith(".xml"):
                continue
            root = etree.fromstring(source.read(name))
            for paragraph in root.xpath(".//w:p[not(.//w:p)]", namespaces=namespace):
                texts = paragraph.xpath(".//w:t", namespaces=namespace)
                if "LiteLLM" not in "".join(node.text or "" for node in texts):
                    continue
                texts[0].text = "{{ short_title }}"
                for node in texts[1:]:
                    node.text = ""
            replacements[name] = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
    _rewrite_members(path, replacements)


def _fix_table_grids(path: Path) -> None:
    """Normalize PDF-converted tables that Word accepts but docxtpl rejects."""
    namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    replacements: dict[str, bytes] = {}
    with ZipFile(path) as source:
        for name in source.namelist():
            if not name.startswith("word/") or not name.endswith(".xml"):
                continue
            payload = source.read(name)
            if b"<w:tbl" not in payload:
                continue
            root = etree.fromstring(payload)
            changed = False
            for table in root.xpath(".//w:tbl", namespaces=namespace):
                if table.find(qn("w:tblGrid")) is not None:
                    continue
                grid = etree.Element(qn("w:tblGrid"))
                cells = table.xpath("./w:tr[1]/w:tc", namespaces=namespace)
                for cell in cells or [None]:
                    column = etree.SubElement(grid, qn("w:gridCol"))
                    widths = cell.xpath("./w:tcPr/w:tcW/@w:w", namespaces=namespace) if cell is not None else []
                    column.set(qn("w:w"), widths[0] if widths else "0")
                table.insert(1, grid)
                changed = True
            if changed:
                replacements[name] = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
    _rewrite_members(path, replacements)


def _metadata(model: DocumentModel) -> dict[str, str]:
    title = model.sections[0].title if model.sections else model.metadata.solution_name
    preamble = "\n".join(item.content for item in model.sections[0].items) if model.sections else ""
    clean = preamble.replace("**", "").replace(">", "")
    version = re.search(r"文档版本[：:]\s*(\S+)", clean)
    date = re.search(r"发布日期[：:]\s*(\d{4}-\d{2}-\d{2})", clean)
    name = re.sub(r"^(部署|Deploying)\s+", "", title).split(" — ", 1)[0].strip()
    name = re.sub(r"\s+(部署指南|Deployment Guide)$", "", name).strip()
    return {
        "document_title": title,
        "short_title": f"{name or model.metadata.project_id} 部署指南",
        "version": version.group(1) if version else model.metadata.version,
        "date": date.group(1) if date else model.metadata.generated_at[:10],
    }


def _body_document(model: DocumentModel, style_source: Path, output: Path) -> None:
    document = Document(style_source)
    body = document._body._element
    final_section = list(body)[-1] if list(body)[-1].tag == qn("w:sectPr") else None
    for element in list(body):
        if element is not final_section:
            body.remove(element)

    with ZipFile(style_source) as archive:
        from xml.etree import ElementTree as ET

        root = ET.fromstring(archive.read("word/document.xml"))
        table_templates = root.findall(".//" + W + "tbl")

    sections = model.sections
    if sections and ("部署指南" in sections[0].title or "Deployment Guide" in sections[0].title):
        sections = sections[0].children

    def append(element) -> None:
        node = parse_xml(ET_to_bytes(element))
        if final_section is None:
            body.append(node)
        else:
            body.insert(len(body) - 1, node)

    def emit(section: Section, level: int = 1) -> None:
        append(_paragraph(section.title, f"Heading{min(level, 5)}"))
        for item in section.items:
            for element in _content_elements(item.content, table_templates):
                append(element)
        for child in section.children:
            emit(child, level + 1)

    for section in sections:
        emit(section)
    document.save(output)

    # The template's ListParagraph style carries an overly broad inherited indent;
    # apply the actual direct paragraph geometry used by its visible list items.
    rendered = Document(output)
    for paragraph in rendered.paragraphs:
        if paragraph.style and paragraph.style.style_id == "ListParagraph":
            indent = paragraph._p.get_or_add_pPr().get_or_add_ind()
            indent.set(qn("w:left"), "1046")
            indent.set(qn("w:right"), "1044")
            indent.set(qn("w:hanging"), "426")
    rendered.save(output)
    _fix_table_grids(output)


def ET_to_bytes(element) -> bytes:
    from xml.etree import ElementTree as ET

    return ET.tostring(element, encoding="utf-8")


def _remove_unused_hyperlinks(path: Path) -> None:
    with ZipFile(path) as source:
        document = source.read("word/document.xml")
        relationships_name = "word/_rels/document.xml.rels"
        if relationships_name not in source.namelist():
            return
        from xml.etree import ElementTree as ET

        relationship_ns = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        used = {
            value
            for node in ET.fromstring(document).iter()
            for attribute, value in node.attrib.items()
            if attribute.startswith(f"{{{relationship_ns}}}")
        }
        relationships = ET.fromstring(source.read(relationships_name))
        for relationship in list(relationships):
            if relationship.get("Type", "").endswith("/hyperlink") and relationship.get("Id") not in used:
                relationships.remove(relationship)
        replacements = {relationships_name: ET.tostring(relationships, encoding="utf-8", xml_declaration=True)}
    _rewrite_members(path, replacements)


def _set_core_properties(path: Path, metadata: dict[str, str]) -> None:
    with ZipFile(path) as source:
        name = "docProps/core.xml"
        payload = source.read(name).replace(b"LiteLLM", metadata["short_title"].encode())
    _rewrite_members(path, {name: payload})


def render_markdown_docxtpl(
    markdown: str | Path,
    template: str | Path,
    style_source: str | Path,
    output: str | Path,
    context_output: str | Path | None = None,
) -> Path:
    model = parse_markdown(markdown)
    metadata = _metadata(model)
    context = {**metadata, "document": model.to_dict()}
    if context_output:
        context_path = Path(context_output)
        context_path.parent.mkdir(parents=True, exist_ok=True)
        context_path.write_text(json.dumps(context, ensure_ascii=False, indent=2), encoding="utf-8")

    output = Path(output)
    output.parent.mkdir(parents=True, exist_ok=True)
    with TemporaryDirectory() as temporary:
        body_path = Path(temporary) / "body.docx"
        _body_document(model, Path(style_source), body_path)
        document = DocxTemplate(template)
        context["body"] = document.new_subdoc(body_path)
        document.render(context, autoescape=True)
        document.save(output)
    _remove_unused_hyperlinks(output)
    _set_core_properties(output, metadata)
    return output
