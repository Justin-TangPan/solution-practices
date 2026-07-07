#!/usr/bin/env python3
"""Generate SAC Technical Report and Business Value Report as DOCX files."""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

BASE = r'C:\Users\Administrator\Desktop\Project\claudeproject\solution-practices'
OUT = os.path.join(BASE, 'reports', 'headroom-claude-code')

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_styled_table(doc, headers, rows):
    """Add a styled table with header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '2F5496')

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, 'D6E4F0')

    return table

def add_code_block(doc, code_text):
    """Add a code block styled paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.0
    # Set shading
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2"/>')
    p._element.get_or_add_pPr().append(shading)
    return p

def gen_technical_report():
    """Generate Technical Implementation Report DOCX."""
    doc = Document()

    # ---- Title Page ----
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    run = p.add_run('技术交付报告')
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Headroom + Claude Code，更省的编程助手')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

    for line in ['', '', '文档类型：SAC 技术交付报告', '方案版本：v1.0', '发布日期：2026-06-09']:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

    # ---- 1. 方案概述 ----
    doc.add_heading('1. 方案概述', level=1)

    doc.add_heading('1.1 方案简介', level=2)
    doc.add_paragraph(
        '本方案面向使用 Claude Code 进行 AI 辅助编程的开发者，通过部署 Headroom 上下文压缩代理，'
        '在华为云 Flexus X 实例上一键搭建"Claude Code + Headroom 代理 + MaaS API"全链路编程环境，'
        '实现 60-95% 的 Token 成本压缩，答案质量不降。'
    )

    doc.add_heading('1.2 应用场景', level=2)
    add_styled_table(doc,
        ['场景', '说明'],
        [
            ['日常开发编程', '开发者使用 Claude Code 进行代码生成、重构、调试，Headroom 自动压缩上下文，降低每轮对话 Token 消耗'],
            ['大规模代码库探索', '在大型项目中搜索和理解代码，工具输出和文件内容可被压缩 80%+，避免上下文溢出中断会话'],
            ['多轮复杂任务', '架构设计、代码审查等多轮交互场景，每次请求携带历史上下文，Headroom 压缩历史日志和工具调用输出'],
            ['日志密集型调试', '排查生产问题时大量日志输出，Headroom 对日志行去重和结构化压缩，压缩率最高达 92%'],
        ]
    )

    # ---- 2. 架构设计 ----
    doc.add_heading('2. 架构设计', level=1)

    doc.add_heading('2.1 技术架构', level=2)
    doc.add_paragraph(
        '本方案采用三层架构：Claude Code CLI（前端交互层）→ Headroom Proxy（成本优化引擎层）→ '
        '华为云 MaaS API（模型推理层）。Headroom 作为中间代理层，透明地拦截 Claude Code 发出的请求，'
        '对上下文做一次性压缩处理后再转发到 MaaS。'
    )

    doc.add_heading('2.2 组件清单', level=2)
    add_styled_table(doc,
        ['组件', '说明', '端口'],
        [
            ['Claude Code CLI', 'Anthropic 官方 CLI 编程工具，通过 ANTHROPIC_BASE_URL 指向本地代理', '—'],
            ['Headroom Proxy', '开源上下文压缩代理，监听 8787 端口，自动压缩后转发到 MaaS', '8787'],
            ['MaaS API 网关', '华为云 ModelArts 模型服务，提供大模型推理 API', '443'],
        ]
    )

    doc.add_heading('2.3 关键配置参数', level=2)
    add_styled_table(doc,
        ['配置项', '值', '作用'],
        [
            ['ANTHROPIC_BASE_URL', 'http://localhost:8787', 'Claude Code → Headroom 代理'],
            ['ANTHROPIC_TARGET_API_URL', 'https://api.modelarts-maas.com/anthropic', 'Headroom 代理 → MaaS 上游'],
            ['ANTHROPIC_AUTH_TOKEN', 'MaaS API Key', '认证凭据'],
            ['ANTHROPIC_MODEL', 'deepseek-v3.2', '使用的模型'],
        ]
    )

    doc.add_heading('2.4 数据流说明', level=2)
    for step in [
        '① 用户在 Terminal 中输入问题 → Claude Code CLI 捕获',
        '② Claude Code 将请求发至 http://localhost:8787（Headroom 代理）',
        '③ Headroom 代理对请求上下文做压缩处理（AST 解析、去重、结构化压缩）',
        '④ 压缩后的请求通过 HTTPS 转发至华为云 MaaS API',
        '⑤ MaaS 调用大模型推理，返回结果',
        '⑥ Headroom 将结果返回给 Claude Code，原始数据本地保存，可逆取回',
    ]:
        doc.add_paragraph(step, style='List Number')

    # ---- 3. 资源规划 ----
    doc.add_heading('3. 资源规划', level=1)

    doc.add_heading('3.1 云资源清单', level=2)
    add_styled_table(doc,
        ['华为云服务', '资源名称', '配置', '数量', '每月预估花费'],
        [
            ['VPC', 'headroom-claude-code-vpc', '172.16.0.0/16', '1', '0.00 元'],
            ['子网', 'headroom-claude-code-subnet', '172.16.1.0/24', '1', '0.00 元'],
            ['安全组', 'headroom-claude-code-sg', 'ICMP、22、8787', '1', '0.00 元'],
            ['Flexus X 实例', 'headroom-claude-code-ecs', '2C4G, Ubuntu 24.04, 40GB SAS', '1', '≈233.60 元'],
            ['弹性公网 IP', 'headroom-claude-code-eip', '动态 BGP, 300Mbit/s', '1', '按流量'],
        ]
    )

    doc.add_heading('3.2 软件依赖清单', level=2)
    add_styled_table(doc,
        ['软件', '版本要求', '安装方式'],
        [
            ['Headroom', 'headroom-ai (pip)', 'pip3 install headroom-ai fastapi uvicorn httpx[http2] transformers'],
            ['Claude Code', '最新版 (npm)', 'npm install -g @anthropic-ai/claude-code'],
            ['Python', '≥ 3.10', '系统预装'],
            ['Node.js', '≥ 18.x', '通过 nvm 安装'],
        ]
    )

    # ---- 4. 部署步骤 ----
    doc.add_heading('4. 部署步骤', level=1)

    doc.add_heading('4.1 准备工作', level=2)
    checks = [
        '已注册华为云账号并完成实名认证',
        '账户不处于欠费或冻结状态',
        '已在 ModelArts 平台开通 MaaS 服务',
        '已获取 MaaS API Key',
    ]
    for c in checks:
        doc.add_paragraph(f'☐ {c}')

    p = doc.add_paragraph()
    run = p.add_run('\nIAM 子用户额外步骤：创建 rf_admin_trust 委托，绑定 Tenant Administrator 策略。')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_heading('4.2 一键部署', level=2)
    doc.add_paragraph(
        '通过华为云解决方案实践搜索"Headroom + Claude Code"或使用 RFS 直接部署。'
        '必填参数为 ECS 密码，部署完成后等待约 10 分钟。'
    )

    doc.add_heading('4.3 部署验证', level=2)
    add_code_block(doc, 'headroom --version')
    add_code_block(doc, 'claude --version')
    add_code_block(doc, 'curl http://localhost:8787/readyz')

    doc.add_heading('4.4 启动 Claude Code', level=2)
    add_code_block(doc,
        'cat >> /root/.bashrc << \'EOF\'\n'
        'export ANTHROPIC_AUTH_TOKEN="your-maas-api-key"\n'
        'export ANTHROPIC_BASE_URL="http://localhost:8787"\n'
        'export ANTHROPIC_TARGET_API_URL="https://api.modelarts-maas.com/anthropic"\n'
        'export ANTHROPIC_MODEL="deepseek-v3.2"\n'
        'EOF\n'
        'source /root/.bashrc\n'
        'claude'
    )

    doc.add_heading('4.5 全链路验证', level=2)
    doc.add_paragraph('在 Claude Code 中发送任意消息，然后检查：')
    add_code_block(doc, 'curl -s http://localhost:8787/stats | grep -o \'"api_requests":[0-9]*\'')
    doc.add_paragraph('预期: api_requests > 0')

    # ---- 5. 验证与调优 ----
    doc.add_heading('5. 验证与调优', level=1)

    doc.add_heading('5.1 功能验证矩阵', level=2)
    add_styled_table(doc,
        ['验证项', '预期结果', '验证方法'],
        [
            ['Headroom 代理运行', '返回 status: healthy', 'curl http://localhost:8787/readyz'],
            ['Claude Code 启动', '进入交互模式', '执行 claude 命令'],
            ['请求经过代理', 'api_requests > 0', 'curl http://localhost:8787/stats'],
            ['Token 压缩生效', '压缩率 60-95%', 'headroom perf'],
            ['可逆取回', '原始内容可恢复', 'MCP 工具 headroom_retrieve'],
        ]
    )

    doc.add_heading('5.2 调优建议', level=2)
    add_styled_table(doc,
        ['优化项', '建议'],
        [
            ['模型选择', '日常用 deepseek-v3.2，复杂任务切 claude-sonnet-4-20250514'],
            ['代理监控', '使用 /metrics 对接 Prometheus，实时观察压缩率和请求量'],
            ['日志管理', 'tail -f /var/log/headroom-proxy.log 实时查看代理运行日志'],
        ]
    )

    # ---- 6. 快速卸载 ----
    doc.add_heading('6. 快速卸载', level=1)
    steps = [
        '登录 RFS 控制台（https://console.huaweicloud.com/rfs/）',
        '找到本方案创建的资源栈，点击右侧"删除"按钮',
        '删除方式选择"删除资源"，输入 Delete，点击"确定"',
    ]
    for s in steps:
        doc.add_paragraph(s, style='List Number')

    # ---- 7. 附录 ----
    doc.add_heading('7. 附录', level=1)
    doc.add_heading('7.1 名词解释', level=2)
    add_styled_table(doc,
        ['术语', '说明'],
        [
            ['Flexus X 实例', '华为云面向中小企业和开发者推出的新一代柔性算力云服务器'],
            ['VPC', '华为云上隔离的、私密的虚拟网络环境'],
            ['EIP', '提供独立的公网 IP 地址和带宽资源'],
            ['Headroom', '开源上下文压缩代理层，减少 LLM Token 消耗'],
            ['Claude Code', 'Anthropic 官方 CLI 编程工具'],
            ['MaaS', 'Model as a Service，华为云模型即服务平台'],
        ]
    )

    # Save
    path = os.path.join(OUT, 'SAC-技术交付报告-Headroom-ClaudeCode.docx')
    doc.save(path)
    print(f'OK: {path}')
    return path


def gen_business_report():
    """Generate Business Value Report DOCX."""
    doc = Document()

    # ---- Title Page ----
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    run = p.add_run('业务价值报告')
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Headroom + Claude Code，更省的编程助手')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

    for line in ['', '', '文档类型：SAC 业务价值报告', '方案版本：v1.0', '发布日期：2026-06-09']:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

    # ---- 1. 产品概述 ----
    doc.add_heading('1. 产品概述', level=1)

    p = doc.add_paragraph()
    run = p.add_run('一句话定位：')
    run.bold = True
    p.add_run('一键部署的 AI 编程辅助环境，通过智能上下文压缩技术，将 Token 成本降低 60-95%，同时保持对话质量和开发体验不变。')

    doc.add_heading('1.1 企业痛点分析', level=2)

    doc.add_heading('挑战一：算力成本失控', level=3)
    add_styled_table(doc,
        ['痛点', '具体表现'],
        [
            ['Token 随对话增长', '每次对话携带完整上下文，代码文件动辄数千行，Token 随对话轮数线性增长'],
            ['大型项目消耗惊人', '单次上下文可超 100K+ Token，月 Token 费用呈指数级攀升'],
            ['成本不可预算', '随着团队使用频率增加，Token 消耗增速超出预期，财务预算难以规划'],
        ]
    )

    doc.add_heading('挑战二：开发效率折损', level=3)
    add_styled_table(doc,
        ['痛点', '具体表现'],
        [
            ['上下文溢出', '复杂对话未完成，上下文窗口已满，被迫中断会话从头开始'],
            ['模型"失忆"', '多轮对话后模型遗忘早期决策上下文，需反复重申背景信息'],
            ['质量下降', '窗口过半后，模型回答精准度明显下降，需人工介入修正'],
        ]
    )

    doc.add_heading('挑战三：部署运维复杂', level=3)
    add_styled_table(doc,
        ['痛点', '具体表现'],
        [
            ['环境配置繁琐', '需自行配置云服务器、网络、安全组、代理层，门槛高'],
            ['缺少标准化方案', '无现成的一键部署方案，每次搭建都是"从零开始"'],
            ['规模化困难', '多团队使用场景下缺少统一管理和监控，推广受阻'],
        ]
    )

    doc.add_heading('1.2 方案价值主张', level=2)
    p = doc.add_paragraph()
    run = p.add_run('核心价值：')
    run.bold = True
    run.font.size = Pt(11)
    doc.add_paragraph(
        'Token 成本 × (1 - 压缩率) + 零中断开发体验 + 一键部署零门槛\n'
        '= 更省（省成本）+ 更省（省时间）+ 更省（省心力）'
    )

    # ---- 2. 业务架构 ----
    doc.add_heading('2. 业务架构图（用户业务流视角）', level=1)

    doc.add_paragraph(
        '本方案从用户业务流视角，分为四层：\n'
        '  • 开发者输入层：通过终端使用 Claude Code，输入自然语言指令\n'
        '  • 成本优化引擎层（核心）：Headroom 自动压缩上下文，不改变开发者使用方式\n'
        '  • 模型推理层：通过华为云 MaaS API 调用大模型\n'
        '  • 企业价值闭环：成本降低 → 效率提升 → 规模化落地 → 持续优化'
    )

    # ---- 3. 应用场景 ----
    doc.add_heading('3. 应用场景（业务视角）', level=1)

    doc.add_heading('场景一：研发团队的日常 AI 编程辅助', level=2)
    doc.add_paragraph('适合客户：中小型研发团队（5-50 人），日常使用 AI 辅助编程')
    doc.add_paragraph('业务痛点：团队 Token 消耗快速增长，月成本不可控；长对话频繁溢出，效率提升打折扣。')
    add_styled_table(doc,
        ['指标', '优化前', '优化后', '降幅'],
        [
            ['月 Token 费（10 人团队）', '¥10,000', '¥1,000-¥4,000', '60-90%'],
            ['对话中断率', '每 3-5 轮中断 1 次', '几乎零中断', '95% ↓'],
            ['新成员接入成本', '30 分钟环境配置', '即刻使用', '零门槛'],
        ]
    )
    p = doc.add_paragraph()
    run = p.add_run('业务收益：')
    run.bold = True
    p.add_run('10 人团队年省 AI 成本 ¥72,000-¥108,000；开发者有效编码时间提升 30%+。')

    doc.add_heading('场景二：大型项目代码维护与重构', level=2)
    doc.add_paragraph('适合客户：中大型企业（50-500 人），管理百万行级代码库')
    add_styled_table(doc,
        ['场景', '压缩效果', '业务影响'],
        [
            ['代码搜索', '17,765 → 1,408 Token（92%）', '探索完整代码库无需中断'],
            ['多轮重构对话', '上下文始终在窗口内', '重构方案一次讨论通过'],
            ['跨模块分析', '可同时加载多个模块文件', '架构师获得全局视野'],
        ]
    )
    p = doc.add_paragraph()
    run = p.add_run('业务收益：')
    run.bold = True
    p.add_run('大型重构任务完成时间缩短 50%，年节省架构师投入约 ¥200,000。')

    doc.add_heading('场景三：日志密集型故障排查', level=2)
    doc.add_paragraph('适合客户：DevOps / SRE 团队，需要 AI 协助排查生产环境问题')
    add_styled_table(doc,
        ['指标', '传统方式', '本方案'],
        [
            ['日志压缩率', '无法完整分析', '92% 压缩，完整保留语义'],
            ['原始日志取回', '需手动翻找', 'headroom_retrieve 一键取回'],
            ['MTTR', '2-4 小时', '1-1.5 小时'],
        ]
    )
    p = doc.add_paragraph()
    run = p.add_run('业务收益：')
    run.bold = True
    p.add_run('MTTR 缩短 40%+，年减少因故障导致的业务损失数百万。')

    # ---- 4. 核心优势 ----
    doc.add_heading('4. 核心优势', level=1)

    doc.add_heading('优势一：行业领先的压缩率 — 60-95% Token 节省', level=2)
    add_styled_table(doc,
        ['工作负载类型', '压缩率'],
        [
            ['代码搜索', '92%'],
            ['文件内容嵌入', '88%'],
            ['MCP 工具输出', '86%'],
            ['日志输出', '92%'],
            ['聊天历史', '89%'],
            ['代码探索', '47%'],
        ]
    )

    doc.add_heading('优势二：零代码改动，透明接入', level=2)
    doc.add_paragraph('无需修改 Claude Code 任何配置或代码。仅需将 ANTHROPIC_BASE_URL 指向 localhost:8787，现有工作流零迁移成本。')

    doc.add_heading('优势三：可逆压缩，不丢信息', level=2)
    doc.add_paragraph('压缩后的原始数据本地保存。模型可通过 MCP 工具 headroom_retrieve 按需取回原始内容。')

    doc.add_heading('优势四：AST 感知的智能压缩', level=2)
    doc.add_paragraph('支持 Python、JS、TS、Go、Rust、Java、C++ 语法树级压缩，保留代码结构和语义。')

    doc.add_heading('优势五：一键部署，即买即用', level=2)
    doc.add_paragraph('基于华为云 RFS 一键部署，10 分钟完成。Headroom + Claude Code 预装，内置 Prometheus 监控。')

    doc.add_heading('优势六：成本和效果可视化', level=2)
    doc.add_paragraph('/stats、/metrics、headroom perf 命令，让每一分钱都花在刀刃上。')

    # ---- 5. 客户收益 ----
    doc.add_heading('5. 客户收益总览', level=1)

    add_styled_table(doc,
        ['收益维度', '量化指标', '业务价值'],
        [
            ['成本节省', 'Token 费用降低 60-95%', '10 人团队年省 ¥72,000-¥108,000'],
            ['效率提升', '有效编码时间 +30%', '团队产出相当于扩编 30% 人力资源'],
            ['部署提效', '一键部署 10 分钟', '从决策到落地，周期从天到小时'],
            ['运维简化', '预装 + 内置监控', '无需专职运维人员'],
            ['弹性扩展', '多种模型切换', '高性价比和高性能灵活选择'],
        ]
    )

    # Save
    path = os.path.join(OUT, 'SAC-业务价值报告-Headroom-ClaudeCode.docx')
    doc.save(path)
    print(f'OK: {path}')
    return path


if __name__ == '__main__':
    os.makedirs(OUT, exist_ok=True)
    t = gen_technical_report()
    b = gen_business_report()
    print(f'\nDone! Generated:')
    print(f'  {t}')
    print(f'  {b}')
