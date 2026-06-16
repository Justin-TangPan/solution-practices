# -*- coding: utf-8 -*-
"""
Markdown to DOCX converter for SAC reports.
Generates professional DOCX files with Chinese font support.
"""

import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os


def create_styled_document():
    """Create a new document with Chinese font support and A4 page size."""
    doc = Document()

    # Set A4 page size
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    # Configure default paragraph style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'SimSun'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

    # Configure heading styles
    for i in range(1, 4):
        heading_style = doc.styles[f'Heading {i}']
        heading_font = heading_style.font
        heading_font.name = 'SimHei'
        heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
        heading_font.color.rgb = RGBColor(0, 0, 0)

        if i == 1:
            heading_font.size = Pt(22)
        elif i == 2:
            heading_font.size = Pt(16)
        elif i == 3:
            heading_font.size = Pt(14)

    # Set default paragraph format for 1.5 line spacing
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 1.5

    return doc


def set_cell_border(cell, **kwargs):
    """Set cell border."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')

    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val["val"]}" '
            f'w:sz="{val["sz"]}" w:space="0" w:color="{val["color"]}"/>'
        )
        tcBorders.append(element)

    tcPr.append(tcBorders)


def add_table_with_borders(doc, headers, rows):
    """Add a table with borders to the document."""
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set table style
    table.style = 'Table Grid'

    # Add headers
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.runs[0]
        run.bold = True
        run.font.name = 'SimHei'
        run.font.size = Pt(11)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')

        # Set header background color
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E6E6E6"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # Add data rows
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = cell_text
            paragraph = cell.paragraphs[0]
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(cell_text)
            run.font.name = 'SimSun'
            run.font.size = Pt(10.5)
            run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

    return table


def add_code_block(doc, code_text):
    """Add a code block with monospace font."""
    # Add a paragraph with monospace font
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.0

    # Set background color for code block
    pPr = paragraph._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:fill="F5F5F5"/>')
    pPr.append(shading)

    # Add border
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="1" w:color="CCCCCC"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="4" w:color="CCCCCC"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="CCCCCC"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="4" w:color="CCCCCC"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    # Set indentation
    ind = parse_xml(f'<w:ind {nsdecls("w")} w:left="360" w:right="360"/>')
    pPr.append(ind)

    # Add code text
    run = paragraph.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
    run.font.color.rgb = RGBColor(0, 0, 0)

    return paragraph


def add_normal_paragraph(doc, text, bold=False):
    """Add a normal paragraph with proper Chinese font."""
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.5

    # Handle bold text
    if bold:
        run = paragraph.add_run(text)
        run.bold = True
        run.font.name = 'SimSun'
        run.font.size = Pt(12)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    else:
        # Check for inline bold
        parts = re.split(r'\*\*(.*?)\*\*', text)
        for i, part in enumerate(parts):
            if part:
                run = paragraph.add_run(part)
                run.font.name = 'SimSun'
                run.font.size = Pt(12)
                run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
                if i % 2 == 1:  # Bold parts
                    run.bold = True

    return paragraph


def add_blockquote(doc, text):
    """Add a blockquote paragraph."""
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.left_indent = Cm(1.0)

    # Set left border
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="12" w:space="4" w:color="4A90D9"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    # Set background color
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:fill="F0F7FF"/>')
    pPr.append(shading)

    run = paragraph.add_run(text)
    run.font.name = 'SimSun'
    run.font.size = Pt(11)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    run.font.color.rgb = RGBColor(51, 51, 51)
    run.italic = True

    return paragraph


def add_bullet_list(doc, items):
    """Add a bullet list."""
    for item in items:
        paragraph = doc.add_paragraph(style='List Bullet')
        paragraph.paragraph_format.line_spacing = 1.5

        # Handle bold text in items
        parts = re.split(r'\*\*(.*?)\*\*', item)
        for i, part in enumerate(parts):
            if part:
                run = paragraph.add_run(part)
                run.font.name = 'SimSun'
                run.font.size = Pt(12)
                run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
                if i % 2 == 1:
                    run.bold = True


def add_numbered_list(doc, items):
    """Add a numbered list."""
    for i, item in enumerate(items, 1):
        paragraph = doc.add_paragraph(style='List Number')
        paragraph.paragraph_format.line_spacing = 1.5

        parts = re.split(r'\*\*(.*?)\*\*', item)
        for j, part in enumerate(parts):
            if part:
                run = paragraph.add_run(part)
                run.font.name = 'SimSun'
                run.font.size = Pt(12)
                run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
                if j % 2 == 1:
                    run.bold = True


def parse_markdown_table(lines, start_idx):
    """Parse a markdown table and return headers, rows, and end index."""
    headers = []
    rows = []
    i = start_idx

    # Parse header
    if i < len(lines) and '|' in lines[i]:
        header_line = lines[i].strip()
        headers = [cell.strip() for cell in header_line.split('|') if cell.strip()]
        i += 1

        # Skip separator line
        if i < len(lines) and '---' in lines[i]:
            i += 1

        # Parse data rows
        while i < len(lines) and '|' in lines[i] and lines[i].strip():
            row_line = lines[i].strip()
            row = [cell.strip() for cell in row_line.split('|') if cell.strip()]
            if row:
                rows.append(row)
            i += 1

    return headers, rows, i


def parse_markdown_to_docx(doc, markdown_content):
    """Parse markdown content and add to document."""
    lines = markdown_content.split('\n')
    i = 0
    in_code_block = False
    code_lines = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Handle code blocks
        if stripped.startswith('```'):
            if in_code_block:
                # End of code block
                code_text = '\n'.join(code_lines)
                if code_text.strip():
                    add_code_block(doc, code_text)
                code_lines = []
                in_code_block = False
            else:
                # Start of code block
                in_code_block = True
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Skip horizontal rules
        if stripped == '---':
            i += 1
            continue

        # Handle headings
        if stripped.startswith('# '):
            heading_text = stripped[2:].strip()
            doc.add_heading(heading_text, level=1)
            i += 1
            continue

        if stripped.startswith('## '):
            heading_text = stripped[3:].strip()
            doc.add_heading(heading_text, level=2)
            i += 1
            continue

        if stripped.startswith('### '):
            heading_text = stripped[4:].strip()
            doc.add_heading(heading_text, level=3)
            i += 1
            continue

        if stripped.startswith('#### '):
            heading_text = stripped[5:].strip()
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.line_spacing = 1.5
            run = paragraph.add_run(heading_text)
            run.bold = True
            run.font.name = 'SimHei'
            run.font.size = Pt(13)
            run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
            i += 1
            continue

        # Handle tables
        if '|' in stripped and not stripped.startswith('-'):
            headers, rows, end_idx = parse_markdown_table(lines, i)
            if headers and rows:
                add_table_with_borders(doc, headers, rows)
                i = end_idx
                continue

        # Handle blockquotes
        if stripped.startswith('>'):
            quote_text = stripped[1:].strip()
            add_blockquote(doc, quote_text)
            i += 1
            continue

        # Handle bullet lists
        if stripped.startswith('- ') or stripped.startswith('* '):
            list_items = []
            while i < len(lines) and (lines[i].strip().startswith('- ') or lines[i].strip().startswith('* ')):
                item_text = lines[i].strip()[2:].strip()
                list_items.append(item_text)
                i += 1
            add_bullet_list(doc, list_items)
            continue

        # Handle numbered lists
        if re.match(r'^\d+\.\s', stripped):
            list_items = []
            while i < len(lines) and re.match(r'^\d+\.\s', lines[i].strip()):
                item_text = re.sub(r'^\d+\.\s', '', lines[i].strip())
                list_items.append(item_text)
                i += 1
            add_numbered_list(doc, list_items)
            continue

        # Handle checklist items
        if stripped.startswith('- [ ]') or stripped.startswith('- [x]'):
            item_text = re.sub(r'^- \[.\]\s*', '', stripped)
            paragraph = doc.add_paragraph(style='List Bullet')
            paragraph.paragraph_format.line_spacing = 1.5
            run = paragraph.add_run(item_text)
            run.font.name = 'SimSun'
            run.font.size = Pt(12)
            run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
            i += 1
            continue

        # Regular paragraph
        add_normal_paragraph(doc, stripped)
        i += 1


def generate_docx(md_file_path, docx_file_path):
    """Generate DOCX from markdown file."""
    print(f"Converting: {md_file_path}")

    # Read markdown content
    with open(md_file_path, 'r', encoding='utf-8') as f:
        markdown_content = f.read()

    # Create styled document
    doc = create_styled_document()

    # Parse markdown and add to document
    parse_markdown_to_docx(doc, markdown_content)

    # Save document
    doc.save(docx_file_path)
    print(f"Generated: {docx_file_path}")


def main():
    """Main function to generate both DOCX files."""
    base_dir = r"C:\Users\Administrator\Desktop\Project\claudeproject\solution-implementations\reports\headroom-opencode"

    # File paths
    md_files = [
        os.path.join(base_dir, "SAC-业务价值报告-Headroom-OpenCode.md"),
        os.path.join(base_dir, "SAC-技术交付报告-Headroom-OpenCode.md")
    ]

    docx_files = [
        os.path.join(base_dir, "SAC-业务价值报告-Headroom-OpenCode.docx"),
        os.path.join(base_dir, "SAC-技术交付报告-Headroom-OpenCode.docx")
    ]

    # Generate DOCX files
    for md_file, docx_file in zip(md_files, docx_files):
        if os.path.exists(md_file):
            generate_docx(md_file, docx_file)
        else:
            print(f"Warning: {md_file} not found")

    print("\nAll DOCX files generated successfully!")


if __name__ == "__main__":
    main()
